# -*- coding: utf-8 -*-
"""Builds RMS_Procurement_Inventory_POS_Operations_Guide.docx — a staff-facing
SOP covering Procurement (PO -> GRC -> GRN), Inventory (central stock, store
stock, transfers, adjustments), and POS/Cashier. Every step, status name, and
"the system will block you" message below is grounded in the actual FastAPI
route logic (purchaseorder_routes.py, grc_routes.py, grn_routes.py,
Stock_transfer_routes.py, inventoryroutes.py, cashier_routes.py) — nothing
here is generic ERP advice invented for the document.
"""
import docx
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Palette (matches the Data Collection & Purpose Register for a
# consistent RMS document family) ──────────────────────────────────────────
NAVY      = RGBColor(0x1E, 0x1B, 0x4B)
ACCENT    = RGBColor(0x4C, 0x1D, 0x95)   # violet-800 — section bands
ACCENT_LT = RGBColor(0xF5, 0xF3, 0xFF)   # violet-50
SLATE     = RGBColor(0x33, 0x41, 0x55)
WHITE     = RGBColor(0xFF, 0xFF, 0xFF)
GREY_LINE = RGBColor(0xD1, 0xD5, 0xDB)
GREEN     = RGBColor(0x0F, 0x76, 0x6E)   # module colour — Procurement
BLUE      = RGBColor(0x1D, 0x4E, 0xD8)   # module colour — Inventory
TEAL      = RGBColor(0x0E, 0x74, 0x90)   # module colour — POS/Cashier
RED       = RGBColor(0x99, 0x1B, 0x1B)   # "system blocks this" marker
AMBER     = RGBColor(0xB4, 0x53, 0x09)   # caution / not-enforced marker
AMBER_BG  = "FEF3C7"                     # caution box fill (amber-100)

MODULE_COLOR = {"Procurement": GREEN, "Inventory": BLUE, "POS / Cashier": TEAL}


def shade_cell(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def shade_paragraph_full_width(paragraph, hex_color):
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    pPr.append(shd)


def box_border(paragraph, hex_color, sz=8):
    """A full paragraph border, used for the caution boxes."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(sz))
        el.set(qn("w:space"), "6")
        el.set(qn("w:color"), hex_color)
        pBdr.append(el)
    pPr.append(pBdr)


doc = docx.Document()

normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(10.5)
normal.font.color.rgb = SLATE

section = doc.sections[0]
section.left_margin = Cm(1.8)
section.right_margin = Cm(1.8)
section.top_margin = Cm(1.6)
section.bottom_margin = Cm(1.6)

# ── Cover page ───────────────────────────────────────────────────────────
for _ in range(4):
    doc.add_paragraph()

eyebrow = doc.add_paragraph()
eyebrow.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = eyebrow.add_run("RETAIL MANAGEMENT SYSTEM  ·  A YARROWTECH PRODUCT")
r.font.size = Pt(11)
r.font.color.rgb = ACCENT
r.bold = True

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("Procurement, Inventory & POS")
r.font.size = Pt(28)
r.bold = True
r.font.color.rgb = NAVY
title2 = doc.add_paragraph()
title2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title2.add_run("Operations Guide")
r.font.size = Pt(28)
r.bold = True
r.font.color.rgb = NAVY

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("What to do, in what order, and what RMS will and will not stop you from doing")
r.font.size = Pt(13)
r.font.color.rgb = SLATE
r.italic = True

for _ in range(2):
    doc.add_paragraph()

meta = doc.add_table(rows=4, cols=2)
meta.alignment = WD_TABLE_ALIGNMENT.CENTER
meta_rows = [
    ("Audience", "Buyers/Merchandisers, Inventory & Warehouse staff, Store Cashiers"),
    ("Covers", "Purchase Orders -> GRC -> GRN  |  Central & Store Inventory  |  POS / Cashier"),
    ("Grounded in", "The actual RMS backend workflow logic — status gates and validation messages, not general retail theory"),
    ("Status", "Working reference — update if the underlying workflow changes"),
]
for row, (k, v) in zip(meta.rows, meta_rows):
    row.cells[0].text = ""
    r = row.cells[0].paragraphs[0].add_run(k)
    r.bold = True
    r.font.color.rgb = ACCENT
    row.cells[1].text = v
    row.cells[0].width = Cm(3.6)
    row.cells[1].width = Cm(11.0)
meta.style = "Table Grid"

doc.add_page_break()

# ── How to read this guide ──────────────────────────────────────────────
h = doc.add_paragraph()
r = h.add_run("How to read this guide")
r.bold = True
r.font.size = Pt(15)
r.font.color.rgb = NAVY

doc.add_paragraph(
    "Each of the three sections below follows the same shape: the real "
    "step-by-step sequence the system enforces, a table of what each status "
    "label means, a list of exactly what RMS will block you from doing (with "
    "the real message you will see), and — where it matters — how a "
    "single-store business differs from a multi-store one."
)
p = doc.add_paragraph()
r = p.add_run("A red marker")
r.bold = True
r.font.color.rgb = RED
p.add_run(" means the system enforces this — it is physically impossible to skip. ")
r2 = p.add_run("An amber caution box")
r2.bold = True
r2.font.color.rgb = AMBER
p.add_run(" means the system will NOT stop you — following the rule is a staff/manager responsibility, not a system guarantee.")

doc.add_paragraph()


# ── Reusable content builders ───────────────────────────────────────────
def add_section_title(number, name):
    band = doc.add_paragraph()
    shade_paragraph_full_width(band, "1E1B4B")
    r = band.add_run(f"  SECTION {number}  ·  {name.upper()}")
    r.bold = True
    r.font.size = Pt(15)
    r.font.color.rgb = WHITE
    band.paragraph_format.line_spacing = 1.9
    band.paragraph_format.space_after = Pt(10)


def add_subheading(text, color=ACCENT):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(12.5)
    r.font.color.rgb = color
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)


def add_steps(steps, color):
    """steps: list of (title, description)"""
    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (title, desc) in enumerate(steps, start=1):
        row = table.add_row().cells
        row[0].width = Cm(1.1)
        row[1].width = Cm(15.3)
        badge = row[0].paragraphs[0]
        badge.alignment = WD_ALIGN_PARAGRAPH.CENTER
        br = badge.add_run(str(i))
        br.bold = True
        br.font.color.rgb = WHITE
        br.font.size = Pt(11)
        shade_cell(row[0], "%02X%02X%02X" % (color[0], color[1], color[2]))
        row[0].vertical_alignment = 1
        p = row[1].paragraphs[0]
        tr = p.add_run(title)
        tr.bold = True
        tr.font.size = Pt(10.5)
        tr.font.color.rgb = SLATE
        p2 = row[1].add_paragraph()
        dr = p2.add_run(desc)
        dr.font.size = Pt(10)
        dr.font.color.rgb = SLATE
    for row in table.rows:
        for cell in row.cells:
            tcPr = cell._tc.get_or_add_tcPr()
            for tag, val in (("w:top", "80"), ("w:bottom", "80"), ("w:left", "80"), ("w:right", "120")):
                pass
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_status_table(rows):
    """rows: list of (status, meaning)"""
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for cell, text, w in zip(hdr, ("Status", "What it means"), (Cm(4.0), Cm(12.4))):
        shade_cell(cell, "EDE9FE")
        run = cell.paragraphs[0].add_run(text)
        run.bold = True
        run.font.color.rgb = ACCENT
        run.font.size = Pt(9.5)
        cell.width = w
    for i, (status, meaning) in enumerate(rows):
        row = table.add_row().cells
        if i % 2 == 1:
            for c in row:
                shade_cell(c, "F5F3FF")
        r0 = row[0].paragraphs[0].add_run(status)
        r0.bold = True
        r0.font.size = Pt(9.5)
        r0.font.name = "Consolas"
        r1 = row[1].paragraphs[0].add_run(meaning)
        r1.font.size = Pt(9.5)
        row[0].width = Cm(4.0)
        row[1].width = Cm(12.4)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_blockers(title, items):
    """items: list of strings — the literal-ish system message plus a short label"""
    p = doc.add_paragraph()
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(10.5)
    r.font.color.rgb = RED
    for label, message in items:
        bp = doc.add_paragraph(style="List Bullet")
        lr = bp.add_run(label + " — ")
        lr.bold = True
        lr.font.size = Pt(10)
        mr = bp.add_run(message)
        mr.font.size = Pt(10)
        mr.italic = True
        mr.font.name = "Consolas"
        mr.font.color.rgb = SLATE
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_caution(title, bullets):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    box_border(p, "F59E0B", sz=10)
    shade_paragraph_full_width(p, AMBER_BG)
    r = p.add_run("⚠  " + title)
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = AMBER
    for b in bullets:
        bp = doc.add_paragraph()
        shade_paragraph_full_width(bp, AMBER_BG)
        box_border(bp, "F59E0B", sz=10)
        br = bp.add_run("•  " + b)
        br.font.size = Pt(10)
        br.font.color.rgb = RGBColor(0x78, 0x35, 0x0F)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def add_note(text):
    p = doc.add_paragraph(text)
    p.runs[0].italic = True
    p.runs[0].font.size = Pt(10)
    p.runs[0].font.color.rgb = SLATE


# ═══════════════════════════════════════════════════════════════════════
# SECTION 1 — PROCUREMENT
# ═══════════════════════════════════════════════════════════════════════
add_section_title(1, "Procurement — PO, GRC & GRN")
doc.add_paragraph(
    "Receiving stock into RMS always follows the same document chain, in "
    "order, with no shortcuts: a Purchase Order must be Approved before a "
    "GRC can be raised against it; a GRC must be Approved before a GRN can "
    "be raised against it; only posting the GRN moves any stock. It is not "
    "possible to receive goods without first raising and approving a GRC."
)

add_subheading("The real sequence", GREEN)
add_steps([
    ("Create the Purchase Order (status: Pending)",
     "Choose a registered, Approved vendor or a walk-in vendor. RMS checks that the vendor has an Approved relationship with your business and that both your KYB and the vendor's KYB are current before allowing the PO."),
    ("Send to Vendor — registered vendors only (status: SentToVendor)",
     "Walk-in vendors do not use this step — share the public PO link with them instead."),
    ("Vendor submits their rates/quantities (status: VendorSubmitted)",
     "RMS compares the vendor's rate to your original line rate. Up to 3% difference is auto-accepted; 3-10% is allowed but flagged; above 10% is blocked until you review and override it."),
    ("Approve the vendor's submission (status: Approved)",
     "This locks in the vendor's rate/quantity as final. RMS auto-reserves matching vendor stock where it can; a shortfall is reported to you but does not block the approval. You may Reject (reason required) instead of approving, or Cancel the PO up to this point."),
    ("Vendor dispatches (delivery stages logged by the vendor)",
     "Production Started -> Ready to Dispatch -> Dispatched. Marking Dispatched requires an expected delivery date, and is the moment reserved vendor stock is actually consumed."),
    ("Goods arrive — raise a GRC (Goods Receipt Challan)",
     "Record what physically arrived (received qty) versus what you're accepting (accepted qty) per line; the difference is automatically logged as a rejection. A GRC starts as Draft."),
    ("Approve the GRC",
     "This is the point of no return for that receipt — an Approved GRC cannot be edited, rejected, or deleted, because it is treated as a legal document. Approving accumulates the received quantity onto the PO and moves the PO to Partially Received or Fully Received."),
    ("Raise a GRN (Goods Receipt Note) against the Approved GRC",
     "A GRN can only reference an Approved GRC, and only one active GRN is allowed per GRC. Line quantities cannot exceed what the GRC accepted."),
    ("Post the GRN",
     "This is the ONLY moment stock actually moves. For a single-store business it posts straight into that store's stock; for a multi-store business it posts into central inventory, which then has to be moved to a specific store via the Inventory module (Section 2). Posting is safe to retry — it will not double-count if run twice."),
], GREEN)

add_subheading("Status meanings", GREEN)
add_status_table([
    ("Pending", "PO created, not yet sent to the vendor."),
    ("SentToVendor", "Visible to the vendor, awaiting their rate/quantity submission."),
    ("VendorSubmitted", "Vendor has responded; awaiting your approval."),
    ("Approved", "You have locked in the vendor's submission."),
    ("WalkinAccepted", "A walk-in (unregistered) vendor has accepted the PO via their share link."),
    ("Rejected", "You declined the vendor's submission (reason required)."),
    ("Cancelled", "PO called off before stock was updated; any reservation/deduction is reversed."),
    ("PartiallyReceived / FullyReceived", "Set automatically as GRCs are approved, based on quantity received so far."),
    ("StockUpdated", "A GRN has been posted against this PO — stock has moved."),
])

add_blockers("What RMS will block you from doing", [
    ("Unapproved vendor", "Vendor 'X' does not have an approved relationship with your retailer..."),
    ("KYB overdue (yours or the vendor's)", "Your retailer business verification (KYB) must be completed before raising new purchase orders..."),
    ("Sending a PO that's already sent", "PO already sent or processed"),
    ("Vendor rate change over 10%", "{n} item(s) have variance > 10%. — you must explicitly override with a reason before the vendor can resubmit."),
    ("Rejecting without a reason, or from the wrong status", "A rejection reason is required. / Only VendorSubmitted POs can be rejected."),
    ("Raising a GRC against the wrong PO status", "PO must be Approved or PartiallyReceived."),
    ("Receiving beyond the PO's allowed tolerance", "Item 'X' exceeds PO tolerance: approved {previous}, this receipt {receiving}, maximum {maximum}. — the allowed over-delivery % is set per PO line and defaults to zero."),
    ("Editing/rejecting/deleting an Approved GRC", "Approved GRCs cannot be edited. They are legal documents."),
    ("Raising a GRN against a GRC that isn't Approved", "GRN can only be created against an Approved GRC."),
    ("A second active GRN against the same GRC", "An active GRN already exists for this GRC. Cancel it before creating a new one."),
    ("GRN quantity exceeding what the GRC accepted", "inwardQty (N) exceeds GRC acceptedQty (M)."),
    ("Editing or deleting a Posted GRN", "Only Draft GRNs can be edited/deleted. Use Cancel for Posted GRNs."),
])

add_note(
    "Single-store vs multi-store: the only branch point in the whole "
    "procurement chain is where a posted GRN lands — directly into your "
    "one store's stock (single-store) or into central inventory, awaiting "
    "transfer to a store (multi-store). Everything upstream of that (PO, "
    "GRC) works identically either way."
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════
# SECTION 2 — INVENTORY
# ═══════════════════════════════════════════════════════════════════════
add_section_title(2, "Inventory — Central Stock, Store Stock & Adjustments")
doc.add_paragraph(
    "A multi-store business has two separate stock pools: central inventory "
    "(what a GRN posts into) and each store's own stock. They only connect "
    "through the Dispatch -> Receive stock-transfer flow below. A "
    "single-store business has no separate central pool at all — GRN "
    "receipts, sales, and adjustments all happen against its one store's "
    "stock directly, and Stock Transfers are unavailable to it entirely."
)

add_subheading("Moving stock from central to a store (multi-store only)", BLUE)
add_note(
    "The old one-step \"allocate\" / \"transfer\" / \"return to central\" "
    "actions no longer exist — using them now returns an error pointing you "
    "to this two-step flow instead."
)
add_steps([
    ("Dispatch (creates a transfer, status: Dispatched)",
     "Stock is deducted from the source immediately — central inventory, or a store's own stock if a store is dispatching. The goods are now \"in transit\": deducted from the source but not yet added anywhere else. A store can only dispatch from itself TO central — not directly to another store."),
    ("Receive (only the destination can do this)",
     "Only stock added at this step reaches the destination — nothing arrives automatically just because it was dispatched. A partial receive is allowed; the remaining quantity stays outstanding on the transfer."),
], BLUE)
add_note(
    "A dispatch can still be edited or cancelled, but only while it is "
    "still Dispatched (not yet received). Once the destination has "
    "received it, the transfer is final."
)

add_subheading("Stock adjustments", BLUE)
doc.add_paragraph(
    "Use the Stock Adjustment screen for any correction to a quantity — "
    "never try to edit the quantity directly on an item's detail page; "
    "that path is blocked precisely so every quantity change carries a "
    "reason and shows up in the Stock Ledger."
)
add_blockers("What RMS will block you from doing", [
    ("Adjustment line with no reason", "Reason is required for SKU 'X'."),
    ("Adjustment that would take stock below zero", "Adjustment would make 'X' stock negative."),
    ("Changing quantity from the item-detail screen", "Use Stock Adjustment to change quantity so the stock ledger remains accurate."),
])
add_caution("Older adjustment route exists but has a gap", [
    "There is a second, older single-line adjustment route still present in the system. It also requires a reason, but it does NOT stop stock from going negative. Always use the standard Stock Adjustment screen, not this legacy route.",
])

add_subheading("Damage & Return", BLUE)
doc.add_paragraph(
    "Damage always deducts stock, and cannot exceed what's actually "
    "available. A Return only puts stock back on the shelf if it's logged "
    "as a Restock — other return actions (e.g. a write-off) leave stock "
    "untouched by design. Neither one requires a separate approval step; "
    "posting it is final immediately."
)

add_subheading("Reorder rules", BLUE)
add_caution("Reorder alerts are not automatic", [
    "A reorder rule (reorder level, reorder quantity, lead time) only marks an item LOW or OK when someone looks at the Low Stock list. Nothing in RMS automatically raises a Purchase Order or pushes an alert when stock crosses the reorder level — checking the list regularly is a manual habit, not a system guarantee.",
])

add_subheading("Single-store vs multi-store, and role differences", BLUE)
doc.add_paragraph(
    "A single-store business gets a few extra capabilities a multi-store "
    "one doesn't have through this module — editing item details directly "
    "and creating size/colour variants — because there is no central "
    "catalogue team to do that separately. A store-scoped login (including "
    "a Store Owner) only ever sees and adjusts its own store's stock; "
    "combined, all-store views are HQ-only."
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════
# SECTION 3 — POS / CASHIER
# ═══════════════════════════════════════════════════════════════════════
add_section_title(3, "POS / Cashier")
doc.add_paragraph(
    "Every cashier session belongs to exactly one store — there is no "
    "\"central\" POS mode. A cashier login only ever sells against, and "
    "affects, its own store's stock."
)

add_subheading("The real sequence", TEAL)
add_steps([
    ("Open the shift",
     "Records the opening cash float. Reopening while a shift is already open just returns the existing one — it will not create a second, duplicate shift."),
    ("Look up / scan the item",
     "RMS checks the live stock quantity for that barcode. An item that has never actually been received into this store via a GRN — for example something only ever seen in a vendor's catalogue — cannot be sold; it is blocked outright."),
    ("Build the bill and save it",
     "Needs at least one valid line (a real barcode, non-zero quantity). Stock is deducted the moment the bill is saved — there is no separate confirm/settle step afterward."),
    ("Process a return, if needed",
     "Uses the same bill screen with the original invoice referenced. The original must exist, be a sale, and be from the same store; each invoice can only be returned once. Returned quantity is put back into store stock automatically."),
    ("Close the shift",
     "Requires an open shift. RMS calculates the expected cash from opening float plus the shift's cash sales, compares it to what the cashier counted, and records any mismatch — as a note for review, not something that blocks closing the till."),
], TEAL)

add_blockers("What RMS will block you from doing", [
    ("Scanning an item never received via GRN", "Product not available for sale — not inducted via GRN."),
    ("Unknown barcode", "No product found for barcode 'X'."),
    ("Saving an empty bill", "No items in bill."),
    ("Return without the original invoice", "Original invoice is required for a return."),
    ("Returning an invoice from another store, or one already returned", "Original invoice was not found in this store. / This invoice has already been returned."),
    ("Closing a shift that was never opened", "No open shift found."),
])

add_caution("Two things this system will NOT stop you from doing", [
    "Overselling: saving a bill for more than the store actually has in stock is allowed — RMS lets the sale go through and only flags a negative-stock warning afterward. It does not refuse the sale the way, say, a stock transfer would.",
    "Uncapped discounts: there is no server-side discount limit or manager-approval check. Whatever discount the POS screen sends is applied as-is. Staying within store discount policy is a cashier/manager responsibility — RMS will not enforce a ceiling for you.",
    "A cash-count mismatch at shift close is logged for review but does not block the shift from closing — it is informational, not a hard stop.",
])

add_subheading("Single-store/multi-store and role differences", TEAL)
doc.add_paragraph(
    "This module doesn't branch by account type the way Procurement and "
    "Inventory do — every cashier, on any tenant, is scoped to exactly one "
    "store. What does differ by role is reporting: an HQ login can see "
    "every store's sales combined or filtered by store; a store-scoped "
    "login only ever sees its own store's numbers, and the cross-store "
    "summary view is HQ-only."
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════
# APPENDIX
# ═══════════════════════════════════════════════════════════════════════
h = doc.add_paragraph()
r = h.add_run("Appendix A — Quick reference")
r.bold = True
r.font.size = Pt(15)
r.font.color.rgb = NAVY

QUICK_REF = [
    ("Procurement", GREEN, [
        ("Do", "Raise a GRC for every physical receipt, even a partial one, before touching stock."),
        ("Do", "Read a >10% rate variance carefully before overriding it — it exists to catch pricing mistakes."),
        ("Don't", "Expect to edit or undo an Approved GRC — it's locked the moment it's approved."),
        ("Don't", "Try to raise a GRN without an Approved GRC — the system won't allow it."),
    ]),
    ("Inventory", BLUE, [
        ("Do", "Always give a real reason on a stock adjustment — it's mandatory and becomes the audit trail."),
        ("Do", "Check the Low Stock / reorder list regularly — nothing pushes that alert to you automatically."),
        ("Don't", "Use the legacy single-line adjustment route — it can silently push stock negative."),
        ("Don't", "Expect a dispatched transfer to arrive on its own — the destination must explicitly Receive it."),
    ]),
    ("POS / Cashier", TEAL, [
        ("Do", "Watch the displayed stock quantity yourself before completing a large sale — the system will let an oversell through."),
        ("Do", "Follow store discount policy manually — RMS applies whatever discount is entered, with no cap."),
        ("Don't", "Assume a cash mismatch at shift close blocks anything — it's logged, not enforced."),
        ("Don't", "Try to sell an item that was never received via GRN — it will not scan as sellable."),
    ]),
]

for name, color, rows in QUICK_REF:
    band = doc.add_paragraph()
    shade_paragraph_full_width(band, "%02X%02X%02X" % (color[0], color[1], color[2]))
    r = band.add_run("  " + name)
    r.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = WHITE
    band.paragraph_format.line_spacing = 1.6
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for cell, text, w in zip(hdr, ("", "Guidance"), (Cm(1.6), Cm(14.8))):
        shade_cell(cell, "F1F5F9")
        run = cell.paragraphs[0].add_run(text)
        run.bold = True
        run.font.size = Pt(9.5)
        cell.width = w
    for i, (tag, text) in enumerate(rows):
        row = table.add_row().cells
        if i % 2 == 1:
            for c in row:
                shade_cell(c, "F8FAFC")
        tr = row[0].paragraphs[0].add_run(tag)
        tr.bold = True
        tr.font.size = Pt(9.5)
        tr.font.color.rgb = GREEN if tag == "Do" else RED
        yr = row[1].paragraphs[0].add_run(text)
        yr.font.size = Pt(9.5)
        row[0].width = Cm(1.6)
        row[1].width = Cm(14.8)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

doc.add_page_break()

h = doc.add_paragraph()
r = h.add_run("Appendix B — Glossary")
r.bold = True
r.font.size = Pt(15)
r.font.color.rgb = NAVY

GLOSSARY = [
    ("PO", "Purchase Order — the buyer's commercial commitment to a vendor."),
    ("GRC", "Goods Receipt Challan — the physical-arrival record: what showed up versus what you're accepting."),
    ("GRN", "Goods Receipt Note — the document that actually posts stock, created against an Approved GRC."),
    ("Tolerance %", "How much over the ordered quantity a vendor is allowed to deliver before RMS blocks the GRC. Set per PO line; zero unless raised."),
    ("Variance %", "How far a vendor's submitted rate differs from the buyer's original PO rate. Above 10% requires an explicit override."),
    ("Central inventory", "The shared stock pool for a multi-store business, fed by GRN, and the source for transfers out to stores."),
    ("Store stock", "A specific store's own stock — what POS actually sells against."),
    ("Dispatch / Receive", "The two-step way stock moves between central and a store, or store to store — deducted on Dispatch, only added on Receive."),
    ("Shift", "A cashier's cash-drawer session, opened with a float and closed with a cash count."),
    ("Walk-in vendor", "A supplier without an RMS account yet, who interacts with a PO through a secure public share link instead of logging in."),
]
table = doc.add_table(rows=1, cols=2)
table.style = "Table Grid"
hdr = table.rows[0].cells
for cell, text, w in zip(hdr, ("Term", "Meaning"), (Cm(4.0), Cm(12.4))):
    shade_cell(cell, "EDE9FE")
    run = cell.paragraphs[0].add_run(text)
    run.bold = True
    run.font.color.rgb = ACCENT
    run.font.size = Pt(9.5)
    cell.width = w
for i, (term, meaning) in enumerate(GLOSSARY):
    row = table.add_row().cells
    if i % 2 == 1:
        for c in row:
            shade_cell(c, "F5F3FF")
    tr = row[0].paragraphs[0].add_run(term)
    tr.bold = True
    tr.font.size = Pt(9.5)
    mr = row[1].paragraphs[0].add_run(meaning)
    mr.font.size = Pt(9.5)
    row[0].width = Cm(4.0)
    row[1].width = Cm(12.4)

doc.add_paragraph()
foot = doc.add_paragraph()
foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = foot.add_run("RMS — a Yarrowtech product  |  Procurement, Inventory & POS Operations Guide  |  Internal working document")
r.italic = True
r.font.size = Pt(9)
r.font.color.rgb = GREY_LINE

out_path = "RMS_Procurement_Inventory_POS_Operations_Guide.docx"
doc.save(out_path)
print("Saved:", out_path)
