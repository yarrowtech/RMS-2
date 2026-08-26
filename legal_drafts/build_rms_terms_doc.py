# -*- coding: utf-8 -*-
"""Builds RMS_Terms_and_Conditions_Revision_Draft.docx — an editable Word
draft mirroring the structure of the reference "Better Pass" clause-revision
table, with content substituted for the RMS platform (Yarrowtech).

NOT LEGAL ADVICE. This is a structured first draft for a lawyer to review
and for Yarrowtech to fill in confirmed details ([TO CONFIRM] placeholders).
"""
import docx
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROWS = [
    ("1", "1.1 Legal Name",
     "Yarrowtech [TO CONFIRM: exact registered legal entity name and "
     "suffix, e.g. “Yarrowtech Private Limited”], operating the "
     "Retail Management System (RMS) platform."),

    ("2", "1.2 Privacy Policy Link",
     "[TO CONFIRM: e.g. https://www.yarrowtech.com/rms/privacy-policy — "
     "insert the actual hosted URL once published.]"),

    ("3", "1.4 Non-acceptance of updated Terms/Privacy Policy",
     "If a Retailer, Vendor, or any Authorized User does not agree to any "
     "updated Terms & Conditions or Privacy Policy, they may discontinue "
     "use of the applicable RMS services and may request account closure. "
     "Continued use of the platform after the effective date of the "
     "updated terms shall constitute acceptance of the revised terms."),

    ("4", "2.1(a)/(b) Vendor & Retailer/Customer definitions",
     "2.1(a) Vendor/Partner:\n"
     "For the purposes of these Terms and Conditions, a “Vendor” or "
     "“Partner” means a supplier, manufacturer, wholesaler, or "
     "service provider registered on the RMS platform to offer products, "
     "raw materials, or job-work services to Retailers, where an Approved "
     "relationship exists between that Vendor and the relevant Retailer. "
     "Vendors/Partners shall include the following categories:\n"
     "1. Product / Finished Goods Vendors – suppliers or manufacturers "
     "offering finished retail products for purchase-order-based sourcing.\n"
     "2. Fabric / Raw Material Suppliers – vendors supplying fabric or "
     "raw materials used in a Retailer’s own production or job-work "
     "operations.\n"
     "3. Job-Work Partners – cutting, stitching, finishing, or other "
     "production-service providers engaged by a Retailer through the "
     "Production & Job Work module.\n\n"
     "2.1(b) Retailer / Tenant (“Customer”):\n"
     "For the purposes of these Terms and Conditions, “Retailer,” "
     "“Tenant,” or “Customer” means a business entity that "
     "registers for and operates an RMS account — including its Head "
     "Office (HQ) Admin, Department Admins, Store Admins, Cashiers, and "
     "other staff acting under that account — to manage inventory, "
     "point-of-sale, purchasing, HR, finance, or vendor relationships "
     "through the RMS platform."),

    ("5", "2.1(d) Retailer/Vendor Data Collected",
     "Clarify exactly what information is collected, by category and "
     "purpose (from the existing RMS product documentation):\n"
     "• Business details (name, GST, address, plan tier) – tenant "
     "identification and invoicing.\n"
     "• Staff identity (name, email, phone, department) – account "
     "access and role assignment.\n"
     "• Vendor profile (catalogue, business type, pricing) – vendor "
     "discovery, approval, and ranking.\n"
     "• Transaction records (sales, purchase orders, GRN) – core "
     "business operations.\n"
     "• Payment outcomes (order IDs, receipts, status) – billing and "
     "access control; RMS never holds raw card/bank/UPI details — those "
     "are handled entirely within Razorpay’s systems.\n"
     "• WhatsApp contact details – order and support notifications, "
     "where opted in.\n"
     "• Usage & audit logs (login activity, key actions) – security "
     "and accountability.\n"
     "• Support ticket content – issue diagnosis and resolution.\n"
     "Refer to the Privacy Policy for processing/storage/retention detail "
     "on each category."),

    ("6", "3.3 Verification Documents (Vendor & Retailer onboarding)",
     "For registration and verification purposes, Yarrowtech may collect "
     "and require Vendors and Retailers to provide information and "
     "documents necessary to establish identity, eligibility, and "
     "legitimacy before an account is activated.\n"
     "(a) Vendors: business/trade name, email address, phone number, "
     "address (city/state), GST or business registration number, website "
     "(if any), a business/catalogue overview, primary verification ID or "
     "business registration document, and bank details for receiving "
     "payment (processed via Razorpay).\n"
     "(b) Retailers: business name, GST/registration details, registered "
     "address, plan tier selected, and the identity details of the HQ "
     "Admin creating the account.\n"
     "Retailer and Vendor onboarding requests are reviewed and approved by "
     "Yarrowtech before an account becomes active. Yarrowtech may request "
     "additional information or documents where reasonably necessary to "
     "verify identity, registration, eligibility, or legitimacy."),

    ("7", "4.3 Refund Policy (subscriptions, not bookings)",
     "All payments made through the RMS platform are processed via "
     "Razorpay and subject to Razorpay’s applicable payment and refund "
     "policies. In addition, the following governs RMS subscription and "
     "billing refunds:\n"
     "(a) Subscription Fees: Plan fees (Basic / Professional / Enterprise) "
     "are billed [TO CONFIRM: monthly/annually] in advance and are "
     "non-refundable for the current billing period once the subscription "
     "is active, except where required by applicable law.\n"
     "(b) Upgrades & Add-ons: Fees for plan upgrades, additional stores, "
     "additional admin seats, or module add-ons (e.g. Production & Job "
     "Work) are billed at purchase and are non-refundable once activated.\n"
     "(c) Vendor Subscription Tiers: Vendor marketplace visibility / "
     "catalogue subscription fees are non-refundable once a tier is "
     "activated for the paid period, except where required by law.\n"
     "(d) Cancellation: A Retailer or Vendor may cancel their subscription "
     "at any time; access continues until the end of the current paid "
     "billing period and will not auto-renew thereafter. [TO CONFIRM: "
     "whether any pro-rata refund applies on annual-plan cancellation.]\n"
     "(e) Processing of Refunds: Any approved refund is processed through "
     "Razorpay to the original payment method, subject to Razorpay’s "
     "processing timelines."),

    ("8", "4.6 Devices / Access",
     "RMS can be accessed via a web browser at [TO CONFIRM: RMS platform "
     "URL]. [If/when a mobile app is released: also via Google Play Store "
     "and Apple App Store on supported devices — remove this sentence if "
     "RMS remains web-only.]"),

    ("10", "5.2",
     "Add the word “correct” before “current,” so the "
     "relevant clause reads “…correct current business and account "
     "information…” Apply to the clause requiring accurate account/"
     "business information (mirrors existing “Accurate onboarding "
     "information” commitment already documented)."),

    ("11", "Privacy Policy – Contractual Agreement",
     "Add a provision making it clear that the Privacy Policy forms part "
     "of the contractual terms governing use of RMS, and that acceptance/"
     "use of the platform constitutes acknowledgement of the Privacy "
     "Policy."),

    ("12", "7.2",
     "This point needs to be reviewed against the actual existing wording "
     "of Clause 7.2 in your base Terms document before filling it in. No "
     "specific change is indicated in the reference material."),

    ("13", "7.4(f)",
     "Add the terms “clone” and “piracy” to the prohibited "
     "activities. The clause should prohibit users from cloning, copying, "
     "reproducing, pirating, reverse engineering, or creating unauthorised "
     "derivative versions of RMS, its platform, source code, content, "
     "systems, or services."),

    ("14", "7.4.10",
     "Add “business data,” “vendor data,” and “transaction "
     "records” to the relevant prohibited-use/data provision. It should "
     "cover unauthorised collection, misuse, disclosure, scraping, or "
     "exploitation of a Retailer’s business data, a Vendor’s catalogue "
     "or pricing data, and transaction records."),

    ("15", "8 – Payments",
     "Payments on RMS are collected for:\n"
     "1. Retailer subscription/plan fees (Basic, Professional, Enterprise)\n"
     "2. Plan upgrades, additional stores, and additional admin seats\n"
     "3. Module add-ons (e.g. Production & Job Work)\n"
     "4. Vendor marketplace subscription / catalogue visibility tiers\n"
     "All payments are processed exclusively through Razorpay; RMS never "
     "collects or stores card, UPI, or bank details directly."),

    ("16", "10.6",
     "Remove this clause, as indicated in the reference material — "
     "confirm against your own base Terms document before removing."),

    ("17", "11.1–11.12 User Responsibilities",
     "11.1 Users are responsible for providing true, accurate, complete, "
     "and current business and identity information when creating an "
     "account, purchasing a subscription, or using any RMS Service.\n"
     "11.2 Users must use RMS only for lawful, genuine retail or vendor "
     "business purposes, in accordance with these Terms and applicable "
     "law.\n"
     "11.3 Users are responsible for reviewing all relevant information "
     "before creating a purchase order, catalogue listing, or vendor "
     "relationship, including pricing, quantities, delivery terms, and "
     "applicable policies.\n"
     "11.4 Users must follow all reasonable instructions, security "
     "requirements, and guidelines provided by Yarrowtech, including "
     "role, department, and store-scope restrictions.\n"
     "11.5 Users are responsible for ensuring their own staff meet any "
     "applicable eligibility or authorization requirements before being "
     "granted access to RMS.\n"
     "11.6 Users must behave professionally and appropriately towards "
     "Yarrowtech, other Retailers, Vendors, and their respective staff. "
     "Users must not engage in abusive, threatening, discriminatory, "
     "fraudulent, or unlawful behaviour.\n"
     "11.7 Users must not misuse RMS, including attempting unauthorised "
     "access, interfering with the platform, introducing malicious code, "
     "creating fraudulent accounts, vendors, or transactions, or using the "
     "Services for any unlawful purpose.\n"
     "11.8 Users are responsible for keeping login credentials "
     "confidential and must report suspected unauthorized access "
     "immediately.\n"
     "11.9 Users are responsible for the accuracy of data they enter "
     "(pricing, stock, tax, vendor details) and for their own regulatory/"
     "tax compliance (e.g. GST); RMS is a record-keeping and operations "
     "tool, not a substitute for compliance.\n"
     "11.10 Users must not provide false, misleading, fraudulent, or "
     "defamatory information, ratings, or feedback concerning Yarrowtech, "
     "RMS, or any Vendor/Retailer.\n"
     "11.11 Users must not transfer, resell, misuse, or commercially "
     "exploit their RMS account, login credentials, or subscription "
     "access except where expressly permitted by Yarrowtech.\n"
     "11.12 Any User who violates these Terms or misuses RMS Services may "
     "have their access restricted, suspended, or terminated, without "
     "prejudice to any other rights or remedies available to Yarrowtech "
     "under these Terms or applicable law."),

    ("18", "15.3 – Trademark",
     "[TO CONFIRM registration status.] If registered: “The RMS name "
     "and logo are trademarks of Yarrowtech.” If not yet registered, "
     "consider filing before asserting trademark rights in the Terms."),

    ("19", "17.3 – Account Deletion / Data Retention",
     "[TO CONFIRM actual practice — do not publish an unconfirmed claim.] "
     "Suggested structure: “Upon deletion of a Retailer or Vendor "
     "account, business data is retained for [TO CONFIRM: X days/years] "
     "for legal, tax, and audit purposes, after which it is permanently "
     "deleted, subject to statutory record-keeping requirements.” Note: "
     "the RMS catalogue module already retains vendor listing media for "
     "30 days after expiry before permanent deletion — confirm whether "
     "that is the intended model for full account deletion, or whether a "
     "different retention period applies."),

    ("20", "Customer Support – New Clause",
     "[TO CONFIRM — insert Yarrowtech’s real support contact details]\n"
     "Mobile: [ ]\n"
     "Email: [ ]\n"
     "Address: [ ]"),
]

doc = docx.Document()

# Base font
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(10.5)

title = doc.add_heading("RMS Terms & Conditions — Clause Revision Draft", level=0)
subtitle = doc.add_paragraph("Retail Management System (RMS) — a Yarrowtech product")
subtitle.runs[0].bold = True
subtitle.runs[0].font.size = Pt(12)

note = doc.add_paragraph()
note_run = note.add_run(
    "DRAFTING AID — NOT LEGAL ADVICE. This table mirrors the structure of "
    "a reference clause-revision document, with content substituted for "
    "RMS/Yarrowtech based on the platform's existing product documentation. "
    "Items marked [TO CONFIRM] are not established facts — verify and "
    "fill them in with real, accurate details, and have this draft "
    "reviewed by a qualified lawyer before publishing it as binding Terms "
    "& Conditions or a Privacy Policy."
)
note_run.italic = True
note_run.font.color.rgb = RGBColor(0x8A, 0x2B, 0x06)

doc.add_paragraph()

table = doc.add_table(rows=1, cols=3)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.style = "Table Grid"

widths = (Cm(1.3), Cm(4.2), Cm(11.5))
hdr = table.rows[0].cells
headers = ("No.", "Clause", "What should be added / changed for RMS")
for cell, text, width in zip(hdr, headers, widths):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = True
    cell.width = width

def set_repeat_header(row):
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    tblHeader = OxmlElement("w:tblHeader")
    tblHeader.set(qn("w:val"), "true")
    trPr.append(tblHeader)

set_repeat_header(table.rows[0])

for no, clause, content in ROWS:
    row = table.add_row().cells
    row[0].text = no
    row[1].text = clause
    row[1].paragraphs[0].runs[0].bold = True
    row[2].paragraphs[0].clear()
    lines = content.split("\n")
    first = True
    for line in lines:
        p = row[2].paragraphs[0] if first else row[2].add_paragraph()
        run = p.add_run(line)
        if line.startswith("[TO CONFIRM"):
            run.italic = True
            run.font.color.rgb = RGBColor(0xB4, 0x53, 0x09)
        first = False
    for cell, width in zip(row, widths):
        cell.width = width

doc.add_page_break()
doc.add_heading("Notes on this draft", level=1)
notes = [
    "Structure mirrors the reference clause-revision table row-for-row "
    "(same numbering, including the gap at row 9, which had no indicated "
    "change in the source material either) so it's easy to cross-check "
    "against the original.",
    "Content for rows 5, 6, 7, 8, 15, and 17 is grounded in RMS's own "
    "existing product/policy documentation (roles, data table, Razorpay-"
    "only payments, subscription/module model) — not invented from the "
    "tourism-platform source.",
    "Rows 1, 2, 8, 18, 19, and 20 contain [TO CONFIRM] placeholders for "
    "facts this draft cannot know: exact registered legal entity name, "
    "Privacy Policy URL, platform access URL, trademark status, data "
    "retention practice on account deletion, and real support contact "
    "details.",
    "Row 7's refund policy is a reasonable SaaS-subscription structure, "
    "not a copy of Better Pass's tour-cancellation-hours model — confirm "
    "billing cycle and whether any pro-rata refund should apply.",
    "Have this reviewed by a lawyer before it is used as binding Terms & "
    "Conditions or a Privacy Policy.",
]
for n in notes:
    doc.add_paragraph(n, style="List Bullet")

out_path = "../legal_drafts/RMS_Terms_and_Conditions_Revision_Draft.docx"
doc.save(out_path)
print("Saved:", out_path)
