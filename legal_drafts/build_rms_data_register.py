# -*- coding: utf-8 -*-
"""Builds RMS_Data_Collection_and_Purpose_Register.docx — a styled,
per-module inventory of what data RMS collects, why, and its purpose
category. Grounded in the actual routes/collections inspected this
session; modules not independently verified are listed but not detailed.
"""
import docx
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Palette ──────────────────────────────────────────────────────────────
NAVY      = RGBColor(0x1E, 0x1B, 0x4B)   # cover / title
ACCENT    = RGBColor(0x4C, 0x1D, 0x95)   # violet-800 — section bands
ACCENT_LT = RGBColor(0xF5, 0xF3, 0xFF)   # violet-50 — light row tint
GOLD      = RGBColor(0xB4, 0x53, 0x09)   # caution / TO-CONFIRM style text
SLATE     = RGBColor(0x33, 0x41, 0x55)
WHITE     = RGBColor(0xFF, 0xFF, 0xFF)
GREY_LINE = RGBColor(0xD1, 0xD5, 0xDB)

PURPOSE_TAGS = {
    "Operational": RGBColor(0x0F, 0x76, 0x6E),
    "Billing":     RGBColor(0xB4, 0x53, 0x09),
    "Security":    RGBColor(0x99, 0x1B, 0x1B),
    "Discovery":   RGBColor(0x1D, 0x4E, 0xD8),
    "Compliance":  RGBColor(0x4C, 0x1D, 0x95),
    "Communication": RGBColor(0x0E, 0x74, 0x90),
}

def shade_cell(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def shade_paragraph_full_width(paragraph, hex_color):
    """Approximates a colored banner by shading the paragraph's own frame."""
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    pPr.append(shd)

doc = docx.Document()

# Base style
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(10.5)
normal.font.color.rgb = SLATE

section = doc.sections[0]
section.left_margin = Cm(1.8)
section.right_margin = Cm(1.8)
section.top_margin = Cm(1.6)
section.bottom_margin = Cm(1.6)

# ── Cover page ───────────────────────────────────────────────────────────
for _ in range(4):
    doc.add_paragraph()

eyebrow = doc.add_paragraph()
eyebrow.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = eyebrow.add_run("RETAIL MANAGEMENT SYSTEM  ·  A YARROWTECH PRODUCT")
r.font.size = Pt(11)
r.font.color.rgb = ACCENT
r.bold = True
r.font.name = "Calibri"

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("Data Collection & Purpose Register")
r.font.size = Pt(30)
r.bold = True
r.font.color.rgb = NAVY

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("What RMS collects, why, module by module")
r.font.size = Pt(14)
r.font.color.rgb = SLATE
r.italic = True

for _ in range(2):
    doc.add_paragraph()

meta = doc.add_table(rows=4, cols=2)
meta.alignment = WD_TABLE_ALIGNMENT.CENTER
meta_rows = [
    ("Document", "Internal data-inventory reference — not a published Privacy Policy"),
    ("Scope", "Modules independently verified in the RMS codebase this review cycle"),
    ("Status", "Working draft — subject to revision as modules change"),
    ("Companion document", "RMS Terms & Conditions — Clause Revision Draft"),
]
for row, (k, v) in zip(meta.rows, meta_rows):
    row.cells[0].text = ""
    r = row.cells[0].paragraphs[0].add_run(k)
    r.bold = True
    r.font.color.rgb = ACCENT
    row.cells[1].text = v
    row.cells[0].width = Cm(4.2)
    row.cells[1].width = Cm(10.5)
meta.style = "Table Grid"

doc.add_page_break()

# ── How to read this ────────────────────────────────────────────────────
h = doc.add_paragraph()
r = h.add_run("How to read this document")
r.bold = True
r.font.size = Pt(15)
r.font.color.rgb = NAVY

intro = (
    "Each module below lists the data fields RMS actually collects at that "
    "point in the platform, why that field is collected, and a purpose "
    "category. Purpose categories: "
)
p = doc.add_paragraph(intro)
for i, (tag, color) in enumerate(PURPOSE_TAGS.items()):
    run = p.add_run(("" if i == 0 else "  ·  ") + tag)
    run.bold = True
    run.font.color.rgb = color

p2 = doc.add_paragraph(
    "This register only details modules whose data model was directly "
    "inspected in the RMS codebase (collections, request/response fields, "
    "and route logic). Modules referenced in RMS's own product documentation "
    "but not independently field-level verified here are listed separately "
    "at the end, rather than guessed at."
)
p2.runs[0].italic = True

doc.add_paragraph()

# ── Module data ──────────────────────────────────────────────────────────
# Each module: (title, one-line description, [(field, why, purpose_tag), ...])
MODULES = [
("1. Platform Account Registration — Retailer / Tenant Onboarding",
 "Collected when a business first signs up for an RMS account and selects a plan.",
 [
    ("Business/company name, registered address, city, state", "Legal identification of the tenant and invoice address", "Compliance"),
    ("GST / business registration number", "Required for tax-compliant invoicing of subscription fees", "Compliance"),
    ("Contact person name, email, phone", "Primary account contact for onboarding and billing communication", "Communication"),
    ("Plan tier selected (Basic / Professional / Enterprise)", "Determines seat, store, and module limits enforced on the account", "Billing"),
    ("Account credentials (email + password/hash)", "Authentication into the RMS platform", "Security"),
 ]),

("2. Admin & Staff Access Management",
 "Collected when an HQ Admin creates or edits staff accounts within their own tenant.",
 [
    ("Staff name, email, phone", "Individual login identity and contact for that staff member", "Security"),
    ("Assigned department(s) — HR, Finance, Inventory, Production & Job Work, etc.", "Drives which modules and routes that login is authorized to reach", "Security"),
    ("Store assignment / scope", "Restricts a store-level login to only that store's data", "Security"),
    ("Role (HQ Admin, Department Admin, Store Admin, Cashier)", "Determines the breadth of actions that login can perform", "Security"),
 ]),

("3. Vendor Onboarding & Identity Verification",
 "Collected when a supplier, fabric vendor, or job-work partner registers on RMS.",
 [
    ("Vendor/business name, email, phone", "Vendor identity for retailer discovery and communication", "Compliance"),
    ("Business type(s) — fabric_supplier, job_worker, general vendor", "Determines which retailer-facing vendor lists the vendor appears in (e.g. Fabric Purchasing vs. Job Work)", "Discovery"),
    ("Address, city, state, pincode", "Vendor location; also the basis for the regional fabric-rate benchmark's state grouping", "Discovery"),
    ("GSTIN / business registration, verification ID", "KYB verification before a retailer relationship can be Approved", "Compliance"),
    ("Bank/payout details (for paid features)", "Processed via Razorpay; RMS itself does not store raw bank/card data", "Billing"),
 ]),

("4. Vendor Catalogue — Finished Goods, Fabric & Job-Work Services",
 "Collected when a vendor lists an item for retailers to discover, inquire on, or directly purchase.",
 [
    ("Item name, category, description, images", "Product presentation to retailers browsing or searching the catalogue", "Discovery"),
    ("Price / price range, MOQ, stock, available sizes/colours, variants", "Buyer decision-making and, where enabled, instant direct-purchase ordering", "Discovery"),
    ("catalogue_kind (finished_goods / fabric_material / job_work_service)", "Routes the listing into the correct buying flow (Quick Order vs. Fabric Purchasing)", "Operational"),
    ("Fabric specs — fabric_type, composition, GSM, width, weave, finish, shade, roll length, rate unit, lead time, testing notes", "Lets a buyer compare fabric on the exact specification that matters (weight, width, composition), and feeds the fabric rate benchmark", "Discovery"),
    ("Job-work specs — service type, rate basis, capacity/day, machine type, accepted materials, lead time", "Lets a buyer evaluate a job-work partner's capacity and service fit", "Discovery"),
    ("Subscription tier at upload, visibility expiry date", "Enforces paid-tier catalogue limits and the visibility window", "Billing"),
 ]),

("5. Inquiries, RFQ & Multi-Vendor Comparison",
 "Collected when a retailer negotiates price/quantity with one or more vendors before committing to a PO.",
 [
    ("Requested size, colour, quantity, target price", "The buyer's specific ask, sent to the vendor for a firm quote", "Operational"),
    ("Buyer note / vendor response note", "Free-text context accompanying the negotiation", "Communication"),
    ("Vendor's confirmed price, quantity, availability", "The vendor's binding response, later convertible into a PO line", "Operational"),
    ("comparison_group_id (multi-vendor RFQ)", "Links several vendors' quotes for the same requirement so they can be compared side by side", "Operational"),
 ]),

("6. Purchase Orders",
 "Collected once a retailer commits to buying from a specific vendor (registered or walk-in).",
 [
    ("Order number, order date, vendor name/ID", "Core transaction identity and audit trail", "Compliance"),
    ("Line items — description, SKU/barcode, quantity, rate, size/colour", "What was actually ordered, at what price", "Operational"),
    ("orderType / purchaseType (e.g. \"Fabric / Raw Material\")", "Routes the PO into the correct downstream pipeline (fabric sheet export, vendor's fabric picker) regardless of which screen created it", "Operational"),
    ("Freight, discount, tax, net amount", "Full landed-cost calculation for the order", "Operational"),
    ("raised_by_user_id / name", "Immutable record of which staff member raised the order", "Security"),
    ("Walk-in vendor share token", "Lets an unregistered vendor view and accept a PO via a secure link without a full account", "Operational"),
 ]),

("7. Fabric Purchasing, Style BOM & Production / Job-Work",
 "Collected when a retailer manufactures finished goods in-house from purchased fabric.",
 [
    ("Style BOM — style name/code, planned quantity, materials, consumption per unit, wastage %", "The \"design\" record: how much fabric a garment is expected to need, used to size the Fabric PO and later checked against actual usage", "Operational"),
    ("Job-work order — job worker/vendor, job type, finished product, expected quantity, due date", "What production work was commissioned and from whom", "Operational"),
    ("Material issue/receipt reconciliation — issued, used, returned, leftover, waste quantities", "Full accountability for fabric handed to a job worker: what was consumed, what came back, what's reusable, what was true loss", "Operational"),
    ("Finished-goods output — barcode, product, quantity, rate", "The manufactured result received back into central inventory", "Operational"),
    ("Vendor acknowledgement — taken date, pieces received, promised-ready date", "Job worker's own confirmation and delivery promise, driving overdue tracking", "Operational"),
    ("Consumption warning (actual vs. BOM-expected usage)", "Flags fabric overuse against the design's expected consumption at receipt time", "Compliance"),
 ]),

("8. Goods Receipt — GRC & GRN",
 "Collected when physical stock (fabric or finished goods) actually arrives at the retailer.",
 [
    ("GRC/GRN number, linked PO number, vendor name", "Links the physical receipt back to the original order", "Compliance"),
    ("Per-item received/accepted/rejected quantity, rejection reason", "Quality and quantity reconciliation against what was ordered", "Operational"),
    ("Barcode (internal), vendor barcode, PO barcode", "Reconciles three different item identifiers — RMS's own, the vendor's label, and the PO's reference — to one stock line", "Operational"),
    ("Fabric spec fields carried from the PO (fabric type, GSM, width, colour, unit, image)", "So a fabric receipt keeps its identity in stock instead of becoming an indistinguishable generic SKU", "Operational"),
    ("Vehicle number, delivery note, received-by", "Physical delivery and accountability record", "Compliance"),
 ]),

("9. Central Inventory & Store Stock",
 "The landing point for everything received via GRN — either central (multi-store) or store-level (single-store tenants).",
 [
    ("Barcode, description, stock quantity, rate/MRP/cost price, unit", "The live, sellable/usable stock record", "Operational"),
    ("is_fabric flag, fabric_type, GSM, width, colour", "Lets fabric stock stay identifiable and searchable after receipt, rather than dissolving into a generic product row", "Operational"),
    ("is_leftover flag, parent barcode", "Marks a reusable fabric remnant as its own pool, distinct from fresh stock, so it can be consciously reused", "Operational"),
    ("Division, section, department, vendor name, source", "Classification and provenance of the stock line", "Operational"),
    ("Adjustment history (qty change, reason, timestamp, source)", "Full audit trail of every stock movement, not just the current balance", "Compliance"),
 ]),

("10. Point of Sale / Cashier & Sales",
 "Collected at the store when a sale is billed.",
 [
    ("Transaction line items, quantities, amounts", "The sale record itself", "Operational"),
    ("Payment method", "Reconciliation of takings by tender type", "Billing"),
    ("Cashier/store identity", "Attributes the sale to a specific till and staff member", "Security"),
 ]),

("11. Stock Adjustments, Transfers, Returns & Damage",
 "Collected whenever stock changes outside of a normal sale or GRN receipt.",
 [
    ("Adjustment quantity and reason", "Explains every non-sale stock movement (correction, shrinkage, damage)", "Compliance"),
    ("Transfer origin store, destination store, quantity", "Tracks stock moving between a retailer's own locations", "Operational"),
    ("Damage/return reason, linked vendor (supplier returns)", "Loss tracking and, where applicable, the vendor-return claim record", "Operational"),
 ]),

("12. Fabric Themes — Multi-Vendor Seasonal Planning",
 "Collected when a buyer groups fabric picks from several vendors under one seasonal requirement.",
 [
    ("Theme name, target date, notes", "The seasonal/collection requirement the picks are being gathered for", "Operational"),
    ("Per-vendor fabric selections (item, spec, quantity, rate) added to the theme", "The buyer's shortlist before finalizing into orders", "Operational"),
    ("Resulting purchase order(s) per vendor on finalize", "One PO per supplier, generated automatically when the theme is finalized — a PO is always single-vendor", "Operational"),
 ]),

("13. Vendor Business Network / B2B Trade",
 "Collected when two vendors trade with each other directly on RMS, separate from retailer purchasing.",
 [
    ("B2B RFQs, orders, receipts, invoices", "Mirrors the retailer PO/GRN flow, but for vendor-to-vendor transactions", "Operational"),
    ("B2B stock and stock ledger entries", "A vendor's own inventory movement record, kept separate from any retailer's stock", "Operational"),
    ("Business connection records", "Tracks which vendors have an approved trading relationship with each other", "Discovery"),
 ]),

("14. Procurement Notifications & Document Conversation",
 "Collected to keep buyers and vendors informed and in direct contact within the context of a specific order.",
 [
    ("Notification type, message, metadata, linked tenant/vendor", "Alerts for events like a new inquiry, PO created, or catalogue item shared", "Communication"),
    ("In-document message text, sender, linked PO/order", "A conversation thread attached to a specific transaction, instead of scattered email/WhatsApp", "Communication"),
 ]),

("15. Subscriptions, Billing & Payments",
 "Collected wherever money moves on RMS — retailer plans, module add-ons, and vendor visibility tiers.",
 [
    ("Retailer plan tier, admin-seat add-ons, store add-ons", "Enforces and bills for the capacity a tenant has actually purchased", "Billing"),
    ("Vendor subscription tier, tier-at-upload", "Enforces catalogue item/photo limits and visibility duration per paid tier", "Billing"),
    ("Razorpay order ID, payment status, receipt", "Confirms paid access; RMS never stores raw card, UPI, or bank details — those stay entirely within Razorpay's systems", "Billing"),
 ]),

("16. Platform Administration & Audit (Super Admin)",
 "Collected at the platform-operator layer, above any individual tenant.",
 [
    ("Audit log entries — action, actor, timestamp", "Accountability and dispute resolution across the whole platform", "Security"),
    ("Onboarding/approval decisions for tenants and vendors", "Record of who was let onto the platform, and when", "Compliance"),
    ("Support ticket content", "Needed to diagnose and resolve the issue the ticket was raised for", "Communication"),
 ]),
]

def add_module(title, desc, rows):
    band = doc.add_paragraph()
    band.paragraph_format.space_before = Pt(4)
    band.paragraph_format.space_after = Pt(2)
    shade_paragraph_full_width(band, "4C1D95")
    r = band.add_run("  " + title)
    r.bold = True
    r.font.size = Pt(13)
    r.font.color.rgb = WHITE
    band.paragraph_format.space_before = Pt(2)
    # padding illusion via leading spaces + line spacing
    band_fmt = band.paragraph_format
    band_fmt.line_spacing = 1.6

    d = doc.add_paragraph()
    dr = d.add_run(desc)
    dr.italic = True
    dr.font.size = Pt(10)
    dr.font.color.rgb = SLATE
    d.paragraph_format.space_after = Pt(6)

    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    widths = (Cm(6.0), Cm(7.5), Cm(2.8))
    hdr = table.rows[0].cells
    for cell, text, w in zip(hdr, ("Data Collected", "Why It's Collected", "Purpose"), widths):
        shade_cell(cell, "EDE9FE")
        p = cell.paragraphs[0]
        run = p.add_run(text)
        run.bold = True
        run.font.color.rgb = ACCENT
        run.font.size = Pt(9.5)
        cell.width = w

    for i, (field, why, tag) in enumerate(rows):
        row = table.add_row().cells
        if i % 2 == 1:
            for c in row:
                shade_cell(c, "F5F3FF")
        row[0].text = ""
        r0 = row[0].paragraphs[0].add_run(field)
        r0.bold = True
        r0.font.size = Pt(9.5)
        row[1].text = ""
        r1 = row[1].paragraphs[0].add_run(why)
        r1.font.size = Pt(9.5)
        row[2].text = ""
        r2 = row[2].paragraphs[0].add_run(tag)
        r2.bold = True
        r2.font.size = Pt(9.5)
        r2.font.color.rgb = PURPOSE_TAGS.get(tag, SLATE)
        for cell, w in zip(row, widths):
            cell.width = w

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

for i, (title, desc, rows) in enumerate(MODULES):
    add_module(title, desc, rows)
    if i < len(MODULES) - 1 and i % 3 == 2:
        doc.add_page_break()

doc.add_page_break()

# ── Not detailed here ───────────────────────────────────────────────────
h = doc.add_paragraph()
r = h.add_run("Modules referenced but not field-level detailed here")
r.bold = True
r.font.size = Pt(15)
r.font.color.rgb = NAVY

p = doc.add_paragraph(
    "RMS's own product documentation lists additional modules — HR, Finance, "
    "Logistics, Design & Pattern, and Forecast & Analytics — that were not "
    "independently inspected at the field level for this register. They are "
    "named here so this document doesn't imply they collect no data; a "
    "follow-up pass should detail them the same way before this register is "
    "treated as complete."
)
for m in ["HR", "Finance", "Logistics", "Design & Pattern", "Forecast & Analytics"]:
    b = doc.add_paragraph(m, style="List Bullet")

doc.add_paragraph()

# ── Data handling principles ────────────────────────────────────────────
h = doc.add_paragraph()
r = h.add_run("Data handling principles observed across every module")
r.bold = True
r.font.size = Pt(15)
r.font.color.rgb = NAVY

principles = [
    "Payment data: RMS never collects or stores raw card, UPI, or bank "
    "details. All payment processing runs through Razorpay; RMS only "
    "retains payment outcomes (order ID, status, receipt).",
    "Access scoping: every field above is only visible to logins whose "
    "role, department, and store scope actually cover it — enforced at "
    "the account level, not left to convention.",
    "Vendor visibility: a vendor's catalogue and pricing are only shown "
    "to a retailer once that retailer has an Approved relationship with "
    "the vendor — never an open, scrapeable marketplace.",
    "Retention on account deletion: not yet established as a confirmed, "
    "documented policy at the time of writing — see the companion Terms "
    "& Conditions draft (clause 17.3) for the open item. The one "
    "confirmed retention rule found in the codebase is narrower: expired "
    "vendor catalogue media is purged 30 days after a listing expires.",
]
for pr in principles:
    doc.add_paragraph(pr, style="List Bullet")

doc.add_paragraph()
foot = doc.add_paragraph()
foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = foot.add_run("RMS — a Yarrowtech product  |  Data Collection & Purpose Register  |  Internal working document")
r.italic = True
r.font.size = Pt(9)
r.font.color.rgb = GREY_LINE

out_path = "../legal_drafts/RMS_Data_Collection_and_Purpose_Register.docx"
doc.save(out_path)
print("Saved:", out_path)
